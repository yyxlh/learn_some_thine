#!/usr/bin/python3
# -*- coding: UTF-8 -*-
import docx

doc=docx.Document('demo.docx')
print(len(doc.paragraphs))

for i in range(0,len(doc.paragraphs)):
    print(doc.paragraphs[i].text)

print(len(doc.paragraphs[1].runs))
for j in range(0,len(doc.paragraphs[1].runs)):
    print(doc.paragraphs[1].runs[j].text)


# 只获取文本
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n\n'.join(fullText)

full_text = getText('demo.docx')
print(full_text)

doc=docx.Document()
doc.add_heading('Header 0',0)
paraobj1=doc.add_paragraph('yangyu,Hello!')
paraobj2=doc.add_paragraph('other line')

# paraobj1.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
paraobj1.add_run('这是后来追加的内容')

doc.add_picture('./out/新年好-五线谱.jpg')
doc.save('hello.docx')